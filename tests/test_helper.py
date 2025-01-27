"""All tests in regards to helper.py."""

import json
import logging
import logging.config
import os
from datetime import datetime
from pathlib import Path

import docx
import docx.table
import pandas as pd
import pytest
import pytz
from churchtools_api.churchtools_api import ChurchToolsApi
from tzlocal import get_localzone

from church_web_helper.helper import (
    extract_relevant_calendar_appointment_shortname,
    get_plan_months_docx,
    get_primary_resource,
    get_special_day_name,
)

logger = logging.getLogger(__name__)

config_file = Path("logging_config.json")
with config_file.open(encoding="utf-8") as f_in:
    logging_config = json.load(f_in)
    log_directory = Path(logging_config["handlers"]["file"]["filename"]).parent
    if not log_directory.exists():
        log_directory.mkdir(parents=True)
    logging.config.dictConfig(config=logging_config)


class Test_Helper:
    """Combined tests"""

    def setup_method(self) -> None:
        """Init API connection used for all tests"""
        self.ct_api = ChurchToolsApi(
            domain=os.getenv("CT_DOMAIN"), ct_token=os.getenv("CT_TOKEN")
        )

    # Parametrized pytest function
    @pytest.mark.parametrize(
        ("sample_input", "expected_output"),
        [
            ("Abendmahl", "mit Abendmahl"),
            ("Familien", "für Familien"),
            ("Grünen", "im Grünen"),
            ("im Sankenbach", "im Grünen"),
            ("auf der Schelkewiese", "im Grünen"),
            ("auf der Gartenschau", "im Grünen"),
            ("Konfirmation", "Konfirmation"),
            ("Goldene Konfirmation", "Konfirmation"),
            ("Ökumenisch", "Ökumenisch"),
            ("Ökum. Gottesdienst", "Ökumenisch"),
            ("Wohnzimmer-Worship", "Wohnzimmer-Worship"),
            ("CVJM-Sonntag", "CVJM"),
            ("Impulsgodi", "Impuls"),
            (
                "10:00 Zentral-Gottesdienst Maki Goldene und Diamantene Konfirmation",
                "Konfirmation",
            ),
        ],
    )
    def test_extract_relevant_calendar_appointment_shortname(
        self, sample_input: str, expected_output: str
    ) -> None:
        """CHeck shortname can be extracted."""
        assert (
            extract_relevant_calendar_appointment_shortname(longname=sample_input)
            == expected_output
        )

    @pytest.mark.parametrize(
        ("date", "expected_output"),
        [
            (
                datetime(year=2024, month=12, day=23, hour=23).astimezone(
                    pytz.timezone("Europe/Berlin")
                ),
                "",
            ),
            (
                datetime(year=2024, month=12, day=24).astimezone(
                    pytz.timezone("Europe/Berlin")
                ),
                "Christvesper",
            ),
            (
                datetime(year=2024, month=12, day=25).astimezone(
                    pytz.timezone("Europe/Berlin")
                ),
                "Christfest I",
            ),
            (
                datetime(year=2024, month=12, day=26).astimezone(
                    pytz.timezone("Europe/Berlin")
                ),
                "Christfest II",
            ),
            (
                datetime(year=2024, month=12, day=26, hour=23).astimezone(
                    pytz.timezone("Europe/Berlin")
                ),
                "Christfest II",
            ),
        ],
    )
    def test_get_special_day_name(self, date: datetime, expected_output: str) -> None:
        """Check that special day names can be identified."""
        special_name_calendar_ids = [52, 72]

        assert (
            get_special_day_name(
                ct_api=self.ct_api,
                special_name_calendar_ids=special_name_calendar_ids,
                date=date,
            )
            == expected_output
        )

    def test_get_plan_months_docx(self) -> None:
        """Check that plan months can be created as docx."""
        df_sample = pd.DataFrame(
            {
                "shortDay": ["3.2", "23.1", "23.1", "3.2"],
                "startDate": [
                    datetime(year=2024, month=2, day=3).astimezone(
                        pytz.timezone("Europe/Berlin")
                    ),
                    datetime(year=2024, month=1, day=23).astimezone(
                        pytz.timezone("Europe/Berlin")
                    ),
                    datetime(year=2024, month=1, day=23).astimezone(
                        pytz.timezone("Europe/Berlin")
                    ),
                    datetime(year=2024, month=2, day=3).astimezone(
                        pytz.timezone("Europe/Berlin")
                    ),
                ],
                "location": ["A1", "A2", "A2", "A3"],
                "shortTime": ["08:00", "10:00", "12:00", "9:00"],
                "predigt": ["P1", "P2", "P1", "P2"],
                "shortName": ["mit Abendmahl", None, None, None],
                "specialService": [None, "mit Kirchenchor", None, None],
                "specialDayName": [None, None, None, None],
            }
        )
        df_sample["startDate"] = df_sample["startDate"].dt.tz_localize(None)
        df_data = (
            df_sample.pivot_table(
                values=["shortTime", "shortName", "predigt", "specialService"],
                index=["specialDayName", "startDate", "shortDay"],
                columns=["location"],
                aggfunc=list,
                fill_value="",
            )
            .reorder_levels([1, 0], axis=1)
            .sort_index(axis=1)
            .reset_index()
            .drop(columns="startDate")
        )

        FILENAME = "tests/samples/test_get_plan_months.docx"
        expected_sample = docx.Document(FILENAME)

        result = get_plan_months_docx(
            df_data,
            from_date=datetime(year=2024, month=1, day=1).astimezone(get_localzone()),
        )

        assert compare_docx_files(result, expected_sample)

    def test_get_primary_resource(self) -> None:
        """Check if primary resource can be identified."""
        SAMPLE_EVENT_ID = 330754
        SAMPLE_DATE = datetime(year=2024, month=9, day=29).astimezone(
            pytz.timezone("Europe/Berlin")
        )
        EXPECTED_RESULT = {"Michaelskirche (MIKI)"}
        RESOURCE_IDS = [8, 16, 17, 20, 21]

        result = get_primary_resource(
            appointment_id=SAMPLE_EVENT_ID,
            relevant_date=SAMPLE_DATE,
            api=self.ct_api,
            considered_resource_ids=RESOURCE_IDS,
        )

        assert result == EXPECTED_RESULT


def test_compare_docx_files() -> None:
    """Check that two docx documents can be compared."""
    FILENAME = "tests/samples/test_get_plan_months.docx"
    FILENAME2 = "tests/samples/test_get_plan_months_other.docx"

    doc1 = docx.Document(FILENAME)
    doc2 = docx.Document(FILENAME)
    assert compare_docx_files(doc1, doc2)[0]

    doc3 = docx.Document(FILENAME2)
    assert not compare_docx_files(doc1, doc3)[0]


def compare_docx_files(
    document1: docx.Document, document2: docx.Document
) -> tuple[bool, str]:
    """Compare both text and table content of two docx files.

    Args:
        document1: _description_
        document2: _description_

    Returns:
        bool - if is equal
        text - description of difference
    """
    # Compare text
    text1 = get_docx_text(document1)
    text2 = get_docx_text(document2)

    if text1 != text2:
        return False, "Text is different"

    # Compare tables
    tables1 = get_docx_tables(document1)
    tables2 = get_docx_tables(document2)

    if not compare_tables(tables1, tables2):
        return False, "Tables are different"

    return True, "Files are identical"


def get_docx_text(document: docx.Document) -> str:
    """Extract text content from the docx file."""
    full_text = [para.text for para in document.paragraphs]
    return "\n".join(full_text)


def get_docx_tables(document: docx.Document) -> str:
    """Extract tables content from the docx file."""
    tables = []

    for table in document.tables:
        table_content = []
        for row in table.rows:
            row_content = [cell.text.strip() for cell in row.cells]
            table_content.append(row_content)
        tables.append(table_content)

    return tables


def compare_tables(tables1: docx.table, tables2: docx.table) -> bool:
    """Compare two sets of tables."""
    if len(tables1) != len(tables2):
        return False

    for table1, table2 in zip(tables1, tables2, strict=False):
        if len(table1) != len(table2):
            return False

        for row1, row2 in zip(table1, table2, strict=False):
            if row1 != row2:
                return False

    return True
