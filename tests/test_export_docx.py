"""All tests in regards to helper.py."""

import json
import logging
import logging.config
import os
from datetime import datetime
from pathlib import Path

import docx
import docx.table
import pandas as pd
import pytz
from churchtools_api.churchtools_api import ChurchToolsApi
from tzlocal import get_localzone

from church_web_helper.export_docx import get_plan_months_docx

logger = logging.getLogger(__name__)

config_file = Path("logging_config.json")
with config_file.open(encoding="utf-8") as f_in:
    logging_config = json.load(f_in)
    log_directory = Path(logging_config["handlers"]["file"]["filename"]).parent
    if not log_directory.exists():
        log_directory.mkdir(parents=True)
    logging.config.dictConfig(config=logging_config)


class Test_Helper:
    """Combined tests."""

    def setup_method(self) -> None:
        """Init API connection used for all tests."""
        self.ct_api = ChurchToolsApi(
            domain=os.getenv("CT_DOMAIN"), ct_token=os.getenv("CT_TOKEN")
        )

    def test_get_plan_months_docx(self) -> None:
        """Check that plan months can be created as docx."""
        df_sample = pd.DataFrame(
            {
                "shortDay": ["3.2", "23.1", "23.1", "3.2"],
                "startDate": [
                    datetime(year=2024, month=2, day=3).astimezone(
                        pytz.timezone("Europe/Berlin")
                    ),
                    datetime(year=2024, month=1, day=23).astimezone(
                        pytz.timezone("Europe/Berlin")
                    ),
                    datetime(year=2024, month=1, day=23).astimezone(
                        pytz.timezone("Europe/Berlin")
                    ),
                    datetime(year=2024, month=2, day=3).astimezone(
                        pytz.timezone("Europe/Berlin")
                    ),
                ],
                "location": ["A1", "A2", "A2", "A3"],
                "shortTime": ["08:00", "10:00", "12:00", "9:00"],
                "predigt": ["P1", "P2", "P1", "P2"],
                "shortName": ["mit Abendmahl", None, None, None],
                "specialService": [None, "mit Kirchenchor", None, None],
                "specialDayName": ["SD2", "SD1", "SD1", "SD2"],
            }
        )
        df_sample["startDate"] = df_sample["startDate"].dt.tz_localize(None)
        df_sample = df_sample.sort_values(
            by=["location", "startDate", "shortDay", "specialDayName", "shortTime"]
        )
        df_data = (
            df_sample.pivot_table(
                values=["shortTime", "shortName", "predigt", "specialService"],
                index=["startDate", "shortDay", "specialDayName"],
                columns=["location"],
                aggfunc=list,
                fill_value="",
            )
            .reorder_levels([1, 0], axis=1)
            .sort_index(axis=1)
            .reset_index()
            .drop(columns="startDate")
        )

        FILENAME = "tests/samples/test_get_plan_months.docx"
        expected_sample = docx.Document(FILENAME)

        result = get_plan_months_docx(
            df_data,
            from_date=datetime(year=2024, month=1, day=1).astimezone(get_localzone()),
        )
        compare_result = compare_docx_files(result, expected_sample)
        assert compare_result[0], compare_result[1]


def test_compare_docx_files():
    """Check that two docx documents can be compared."""
    FILENAME = "tests/samples/test_get_plan_months.docx"
    FILENAME2 = "tests/samples/test_get_plan_months_other.docx"

    doc1 = docx.Document(FILENAME)
    doc2 = docx.Document(FILENAME)
    assert compare_docx_files(doc1, doc2)[0]

    doc3 = docx.Document(FILENAME2)
    assert not compare_docx_files(doc1, doc3)[0]

# Following methods assist with test cases

def compare_docx_files(
    document1: docx.Document, document2: docx.Document
) -> tuple[bool, str]:
    """Compare both text and table content of two docx files.

    Args:
        document1: _description_
        document2: _description_

    Returns:
        bool - if is equal
        text - description of difference
    """
    # Compare text
    text1 = get_docx_text(document1)
    text2 = get_docx_text(document2)

    if text1 != text2:
        return False, "Text is different"

    # Compare tables
    tables1 = get_docx_tables(document1)
    tables2 = get_docx_tables(document2)

    table_compare_result = compare_tables(tables1, tables2)
    if not table_compare_result[0]:
        return table_compare_result[0], table_compare_result[1]

    return True, "Files are identical"


def get_docx_text(document: docx.Document) -> str:
    """Extract text content from the docx file."""
    full_text = [para.text for para in document.paragraphs]
    return "\n".join(full_text)


def get_docx_tables(document: docx.Document) -> str:
    """Extract tables content from the docx file."""
    tables = []

    for table in document.tables:
        table_content = []
        for row in table.rows:
            row_content = [cell.text.strip() for cell in row.cells]
            table_content.append(row_content)
        tables.append(table_content)

    return tables


def compare_tables(tables1: docx.table, tables2: docx.table) -> (bool, str):
    """Compare two sets of tables.

    Returns: if identical, and reason
    """
    if len(tables1) != len(tables2):
        return False, "length of table does not match"

    for table1, table2 in zip(tables1, tables2, strict=True):
        if len(table1) != len(table2):
            return False, "length of tables elements does not match"

        for row1, row2 in zip(table1, table2, strict=True):
            if row1 != row2:
                return False, f"row is different: {row1}, {row2}"

    return True, "identical"
