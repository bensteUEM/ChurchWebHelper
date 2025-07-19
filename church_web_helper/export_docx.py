"""This module implements all helper functions specific to docx export."""

import locale
import logging
from datetime import datetime

import docx
import docx.table
import pandas as pd
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)


def get_plan_months_docx(
    data: pd.DataFrame, from_date: datetime, apply_locale: str = "de_DE.UTF-8"
) -> docx.Document:
    """Function which converts a Dataframe into a DOCx document.

    used for final print modifications.

    Args:
        data: pre-formatted data to be used as base
        from_date: date used for heading
        apply_locale: the locale to be used (in particular for date to text conversion)
            e.g. "de_DE.UTF-8" (default) or "deu" (for windows)

    Returns:
        document reference
    """
    locale.setlocale(locale.LC_TIME, apply_locale)

    document = docx.Document()
    padding_left = 1.5
    padding_right = -0.25
    padding_top = -1
    set_page_margins(
        document,
        top=5.71 + padding_top,
        bottom=1.27,
        left=2.75 - padding_left,
        right=0.25 - padding_right,
    )

    heading = f"Unsere Gottesdienste im {from_date.strftime('%B %Y')}"
    paragraph = document.add_heading(heading)
    for run in paragraph.runs:
        run.bold = True
        run.font.name = "ArialNarrow"
        run.font.size = Pt(32)
        run.font.color.rgb = RGBColor.from_string("000000")

    locations = list(dict.fromkeys(item[0] for item in data.columns[2:]))

    table = document.add_table(rows=1, cols=len(locations) + 1)
    hdr_cells = table.rows[0].cells

    for column_no, content in enumerate(locations):
        hdr_cells[column_no + 1].text = content
        for paragraph in hdr_cells[column_no + 1].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for _index, df_row in data.iterrows():
        row_cells = table.add_row().cells
        para = row_cells[0].paragraphs[0]
        para.add_run(df_row["shortDay"].iloc[0]).add_break()
        para.add_run(df_row["specialDayName"].iloc[0])
        for paragraph in row_cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        for column_no, location in enumerate(locations):
            generate_event_paragraph(
                target_cell=row_cells[1 + column_no], relevant_entry=df_row[location]
            )

    change_table_format(table=table)

    FOOTER_TEXTs = [  # noqa: N806
        "Sonntags um 10.00 Uhr findet regelmäßig Kinderkirche in Baiersbronn statt. "
        "Bei Interesse melden Sie sich bitte direkt bei den Mitarbeitenden.: "
        "Juliane Haas, Tel: 604467",
        "Aktuelle und weitere Termine auch auf unserer Website",
    ]
    for footer_text in FOOTER_TEXTs:
        para = document.add_paragraph(footer_text)
        run = para.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(11)
    logger.info("Finished get_plan_months_docx")
    return document


def generate_event_paragraph(
    target_cell: docx.table._Cell, relevant_entry: pd.Series
) -> None:
    """Function which generates the content of one table cell.

    Used with get_plan_months_docx
    Iterates through all items in relevant row
    and using the columns to generate the text.

    Args:
        target_cell: the table cell which should get the content
        relevant_entry: the pd series with list of items in each column

    Returns:
        None because working inplace
    """
    for entry_index in range(1 - 1, len(relevant_entry["shortTime"])):
        current_paragraph = (
            target_cell.paragraphs[0]
            if entry_index == 0
            else target_cell.add_paragraph("")
        )
        if relevant_entry["shortTime"][entry_index]:
            current_paragraph.add_run(relevant_entry["shortTime"][entry_index])
        if relevant_entry["shortName"][entry_index]:
            # should be single relevant_row only but getting list
            current_paragraph.add_run(" " + relevant_entry["shortName"][entry_index])

        # Apply bold formatting nad set font size and font family
        for run in current_paragraph.runs:
            run.bold = True

        if relevant_entry["specialService"][entry_index]:
            current_paragraph.runs[-1].add_break()
            current_paragraph.add_run(
                " " + relevant_entry["specialService"][entry_index]
            )
        if relevant_entry["predigt"][entry_index]:
            current_paragraph.runs[-1].add_break()
            current_paragraph.add_run(f"({relevant_entry['predigt'][entry_index]})")


def change_table_format(table: docx.table) -> None:
    """Inplace overwrite of styles.

    Args:
        table: the table to modify
    """
    # Access the XML element of the table and move ident
    # because by default it's 1,9cm off
    tbl_pr = table._element.xpath("w:tblPr")[0]  # noqa: SLF001
    tbl_indent = OxmlElement("w:tblInd")
    tbl_indent.set(qn("w:w"), "107.12")
    tbl_indent.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_indent)

    # iterate all rows
    for row in table.rows:
        # iterate all cells
        for cell in row.cells:
            set_cell_border(cell=cell)
            set_cell_margins(cell, 100, 100, 0, 100)
            # iterate all paragraphs
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(100) * 20
                for run in paragraph.runs:
                    run.font.name = "ArialNarrow"
                    run.font.size = Pt(15)


def set_page_margins(
    doc: docx.Document, top: float, bottom: float, left: float, right: float
) -> None:
    """Helper to set document page borders in cm.

    Args:
        doc: the document to change
        top: border in cm
        bottom: border in cmon_
        left: border in cm
        right: border in cm
    """
    section = doc.sections[0]

    # Set the margins
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def set_cell_border(cell) -> None:
    """Function to add borders to a cell.

    Args:
        cell: the table cell to change
    """
    tc = cell._element  # noqa: SLF001
    tcPr = tc.get_or_add_tcPr()

    # Create borders element
    tcBorders = OxmlElement("w:tcBorders")

    # Define each side's border attributes (top, left, bottom, right)
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")  # Border style
        border.set(qn("w:sz"), "4")  # Border width (in eighths of a point)
        border.set(qn("w:color"), "auto")  # Automatic color (black)
        tcBorders.append(border)

    tcPr.append(tcBorders)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0) -> None:
    """Function to set cell margins (padding).

    Args:
        cell: the table cell to modify
        top: margin size in dxa (pt/20). Defaults to 0.
        start: margin size in dxa (pt/20). Defaults to 0.
        bottom: margin size in dxa (pt/20). Defaults to 0.
        end: margin size in dxa (pt/20). Defaults to 0.
    """
    tc = cell._element  # Access the underlying XML element for the cell  # noqa: SLF001
    tcPr = tc.find(ns.qn("w:tcPr"))  # Find the <w:tcPr> element if it exists

    # If <w:tcPr> doesn't exist, create it
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    # Create or modify the tcMar (cell margins) element
    tcMar = tcPr.find(ns.qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    # Set each margin (top, start=left, bottom, end=right) in dxa (1/20th of a point)
    for side, margin in [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ]:
        margin_element = OxmlElement(f"w:{side}")
        margin_element.set(ns.qn("w:w"), str(margin))  # Set margin size in dxa
        margin_element.set(ns.qn("w:type"), "dxa")  # dxa = 1/20th of a point
        tcMar.append(margin_element)
