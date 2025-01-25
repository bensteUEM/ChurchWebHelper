"""File which includes helpful methods for data transformation in the context of ChurchWebHelper.

It is used to outsource parts which don't need to be part of app.py
"""

import logging
from collections import OrderedDict
from datetime import datetime

import docx
import docx.table
import pandas as pd
import xlsxwriter
from churchtools_api.churchtools_api import ChurchToolsApi as CTAPI
from dateutil.relativedelta import relativedelta
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)


def get_special_day_name(
    ct_api: CTAPI, special_name_calendar_ids: list[int], date: datetime
) -> str:
    """Retrieve the name of the first calendarf entry one a specific day.

    This can be used to retrieve "holiday" names if they are specfied in the calendar

    Args:
        ct_api: initialized churchtools api connection which should be used as datasource
        special_name_calendar_ids: list of calendar ids used for special names
        date: the day used for lookup

    Returns:
        str: first result of calendar name - usually name of a holiday
    """
    trunc_date = datetime(year=date.year, month=date.month, day=date.day)

    special_names = ct_api.get_calendar_appointments(
        calendar_ids=special_name_calendar_ids,
        from_=trunc_date,
        to_=trunc_date + relativedelta(days=1) - relativedelta(seconds=1),
    )
    if special_names and len(special_names) > 0:
        return special_names[0]["caption"]

    return ""


def extract_relevant_calendar_appointment_shortname(longname: str) -> str:
    """Tries to extract a shortname for "special ocasion events" based on a mapping.

    Excecution order is relevant!

    Args:
        longname (str): the original title of the event

    Returns:
        str: the shortened version which only includes special parts. Does return empty str if nothing detected
    """
    longname = longname.upper()
    known_keywords = ["Abendmahl", "Familien", "Grünen", "Konfirmation", "Ökum"]
    for keyword in known_keywords:
        if keyword.upper() in longname:
            result = keyword
            break
    else:
        known_specials = {
            "Sankenbach": "Grünen",
            "Flößerplatz": "Grünen",
            "Gartenschau": "Grünen",
            "Schelkewiese": "Grünen",
            "Wohnzimmer": "Wohnzimmer",
            "CVJM": "CVJM",
            "Impuls": "Impuls",
        }
        for keyword, replacement in known_specials.items():
            if keyword.upper() in longname:
                result = replacement
                break
        else:
            result = ""

    fulltext_replacements = {
        "Abendmahl": "mit Abendmahl",
        "Familien": "für Familien",
        "Grünen": "im Grünen",
        "Wohnzimmer": "Wohnzimmer-Worship",
        "Ökum": "Ökumenisch",
    }
    for old, new in fulltext_replacements.items():
        if result == old:
            result = result.replace(old, new)
            break

    return result


def get_plan_months_docx(data: pd.DataFrame, from_date: datetime) -> docx.Document:
    """Function which converts a Dataframe into a DOCx document used for final print modifications.

    Args:
        data: pre-formatted data to be used as base
        from_date: date used for heading

    Returns:
        document reference
    """
    document = docx.Document()
    padding_left = 1.5
    padding_right = -0.25
    padding_top = -1
    set_page_margins(
        document,
        top=5.71 + padding_top,
        bottom=1.27,
        left=2.75 - padding_left,
        right=0.25 - padding_right,
    )

    heading = f"Unsere Gottesdienste im {from_date.strftime('%B %Y')}"
    paragraph = document.add_heading(heading)
    for run in paragraph.runs:
        run.bold = True
        run.font.name = "ArialNarrow"
        run.font.size = Pt(32)
        run.font.color.rgb = RGBColor.from_string("000000")

    locations = {item[0] for item in data.columns[2:]}

    table = document.add_table(rows=1, cols=len(locations) + 1)
    hdr_cells = table.rows[0].cells

    for column_no, content in enumerate(locations):
        hdr_cells[column_no + 1].text = content
        for paragraph in hdr_cells[column_no + 1].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for _index, df_row in data.iterrows():
        row_cells = table.add_row().cells
        para = row_cells[0].paragraphs[0]
        para.add_run(df_row["shortDay"].iloc[0]).add_break()
        para.add_run(df_row["specialDayName"].iloc[0])
        for paragraph in row_cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        for column_no, location in enumerate(locations):
            generate_event_paragraph(
                target_cell=row_cells[1 + column_no], relevant_entry=df_row[location]
            )

    change_table_format(table=table)

    FOOTER_TEXTs = [
        "Sonntags um 10.00 Uhr findet regelmäßig Kinderkirche in Baiersbronn statt. Bei Interesse melden Sie sich bitte direkt bei den Mitarbeitenden.: Juliane Haas, Tel: 604467 oder Bärbel Vögele, Tel.:121136",
        "Aktuelle und weitere Termine auch auf unserer Website",
    ]
    for footer_text in FOOTER_TEXTs:
        para = document.add_paragraph(footer_text)
        run = para.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(11)
    logger.info("Finished get_plan_months_docx")
    return document


def get_plan_months_xlsx(
    data: pd.DataFrame, from_date: datetime, filename: str
) -> xlsxwriter.Workbook:
    """Function which converts a Dataframe into a XLXs used as admin overview printout.

    Args:
        data: pre-formatted data to be used as base
        from_date: date used for heading
        filename: name of the file including extension

    Returns:
        workbook reference
    """
    workbook = xlsxwriter.Workbook(filename)

    heading = f"{from_date.strftime('%B %Y')}"
    worksheet = workbook.add_worksheet(name=heading)

    locations = {item[0] for item in data.columns[2:]}

    row = 0

    # setup 3 header lines
    NUMBER_OF_COLUMNS_PER_LOCATION = 4
    format_header = workbook.add_format({"bold": True})
    format_header_b = workbook.add_format({"bold": True, "bottom": 2})
    format_header_bl = workbook.add_format({"bold": True, "bottom": 2, "left": 2})
    format_header_l = workbook.add_format({"bold": True, "left": 2})

    for location_index, location_name in enumerate(locations):
        worksheet.write(
            row,
            location_index * NUMBER_OF_COLUMNS_PER_LOCATION + 1,
            location_name,
            format_header_l,
        )
        for offset, header in enumerate(["Uhr-", "Prediger", "Abm", "Organist"]):
            worksheet.write(
                row + 1,
                1 + location_index * NUMBER_OF_COLUMNS_PER_LOCATION + offset,
                header,
                format_header_l if offset == 0 else format_header,
            )
        for offset, header in enumerate(["zeit", "", "Taufe", "Musik"]):
            worksheet.write(
                row + 2,
                1 + location_index * NUMBER_OF_COLUMNS_PER_LOCATION + offset,
                header,
                format_header_bl if offset == 0 else format_header_b,
            )

    row += 3

    """
    Each location should have a 2*4 entry by eventwhich looks like this

    TIME      | Predigt | ABM   | Organist
    shortName |         | Taufe | Musik

    """

    column_offsets = [0, 1, 2, 3, 0, 1, 2, 3]
    row_offsets = [0, 0, 0, 0, 1, 1, 1, 1]
    column_references = [
        "shortTime",
        "predigt_lastname",
        "abendmahl",
        "organist_lastname",
        "shortName",
        None,
        "taufe",
        "musik",
    ]

    # Iterate all data rows
    location_column_offset = 0
    for _index, df_row in data.iterrows():
        format_content = workbook.add_format({"bold": True, "border": 1})
        format_content_b = workbook.add_format({"bold": True, "border": 1, "bottom": 2})

        worksheet.write(row, 0, df_row["shortDay"].iloc[0], format_content)
        worksheet.write(row + 1, 0, df_row["specialDayName"].iloc[0], format_content_b)

        max_events_per_date = max([len(i) for i in df_row[slice(None), "shortTime"]])
        for row_offset, column_offset, column_value in zip(
            row_offsets, column_offsets, column_references, strict=False
        ):
            for location_index, location in enumerate(locations):
                location_column_offset = location_index * NUMBER_OF_COLUMNS_PER_LOCATION

                for event_per_day_offset in range(max_events_per_date):
                    value = ""
                    if column_value in df_row[location].index:
                        if len(df_row[location, column_value]) > event_per_day_offset:
                            value = df_row[location, column_value][event_per_day_offset]

                    cell_format = workbook.add_format({"border": 1})
                    if row_offset == 1:
                        cell_format.set_bottom(2)
                    if column_offset == 0:
                        cell_format.set_left(2)

                    worksheet.write(
                        row + row_offset + event_per_day_offset * 2,
                        1 + location_column_offset + column_offset,
                        str(value),
                        cell_format,
                    )

        row += 2 * max_events_per_date
    logger.info("Finished get_plan_months_xlsx")
    return workbook


def deduplicate_df_index_with_lists(df_input: pd.DataFrame) -> pd.DataFrame:
    """Flattens a df with multiple same index entries to list entries.

    Args:
        df_input: the original dataframe which contains multiple entries for "shortDay" index per column

    Returns:
        flattened df which has unique shortDay and combined lists in cells
    """
    shortDays = list(OrderedDict.fromkeys(df_input["shortDay"]).keys())
    df_output = pd.DataFrame(columns=df_input.columns)
    for shortDay in shortDays:
        df_shortDay = df_input[df_input["shortDay"] == shortDay]
        new_index = len(df_output)
        df_output.loc[new_index] = [pd.NA] * df_output.shape[1]
        df_output.loc[new_index, "shortDay"] = df_shortDay["shortDay"].iloc[0]
        df_output.loc[new_index, "specialDayName"] = df_shortDay["specialDayName"].iloc[
            0
        ]

        locations = OrderedDict.fromkeys(i[0] for i in df_shortDay.columns[2:])
        for location in locations:
            for col in df_shortDay[location]:
                value_list = []
                df_non_empty = df_shortDay[location][
                    ~(
                        df_shortDay[location].apply(
                            lambda row: (row == "").all(), axis=1
                        )
                    )
                ]
                for value in df_non_empty[col]:
                    if isinstance(value, list):
                        value_list.extend(value)
                    else:
                        value_list.append(value)
                df_output.loc[new_index, (location, col)] = value_list

    df_output = df_output.fillna("")
    logger.debug("finished deduplicate_df_index_with_lists")

    return df_output


def generate_event_paragraph(
    target_cell: docx.table._Cell, relevant_entry: pd.Series
) -> None:
    """Function which generates the content of one table cell.

    Used with get_plan_months_docx
    Iterates through all items in relevant row and using the columns to generate the text.

    Args:
        target_cell: the table cell which should get the content
        relevant_row: the pd series with list of items in each column

    Returns:
        None because working inplace
    """
    for entry_index in range(1 - 1, len(relevant_entry["shortTime"])):
        current_paragraph = (
            target_cell.paragraphs[0]
            if entry_index == 0
            else target_cell.add_paragraph("")
        )
        if relevant_entry["shortTime"][entry_index]:
            # TODO #16 something is off here ... same time for all on web export
            current_paragraph.add_run(relevant_entry["shortTime"][entry_index])
        if relevant_entry["shortName"][entry_index]:
            # should be single relevant_row only but getting list
            current_paragraph.add_run(" " + relevant_entry["shortName"][entry_index])

        # Apply bold formatting nad set font size and font family
        for run in current_paragraph.runs:
            run.bold = True

        if relevant_entry["specialService"][entry_index]:
            current_paragraph.runs[-1].add_break()
            current_paragraph.add_run(
                " " + relevant_entry["specialService"][entry_index]
            )
        if relevant_entry["predigt"][entry_index]:
            current_paragraph.runs[-1].add_break()
            current_paragraph.add_run(f"({relevant_entry['predigt'][entry_index]})")


def change_table_format(table: docx.table) -> None:
    """Inplace overwrite of styles.

    Args:
        table: the table to modify
    """
    # Access the XML element of the table and move ident because by default it's 1,9cm off
    tbl_pr = table._element.xpath("w:tblPr")[0]
    tbl_indent = OxmlElement("w:tblInd")
    tbl_indent.set(qn("w:w"), "107.12")
    tbl_indent.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_indent)

    # iterate all rows
    for row in table.rows:
        # iterate all cells
        for cell in row.cells:
            set_cell_border(cell=cell)
            set_cell_margins(cell, 100, 100, 0, 100)
            # iterate all paragraphs
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(100) * 20
                for run in paragraph.runs:
                    run.font.name = "ArialNarrow"
                    run.font.size = Pt(15)


def set_page_margins(
    doc: docx.Document, top: float, bottom: float, left: float, right: float
) -> None:
    """Helper to set document page borders in cm.

    Args:
        doc: the document to change
        top: border in cm
        bottom: border in cmon_
        left: border in cm
        right: border in cm
    """
    section = doc.sections[0]

    # Set the margins
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def set_cell_border(cell) -> None:
    """Function to add borders to a cell.

    Args:
        cell: the table cell to change
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Create borders element
    tcBorders = OxmlElement("w:tcBorders")

    # Define each side's border attributes (top, left, bottom, right)
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")  # Border style
        border.set(qn("w:sz"), "4")  # Border width (in eighths of a point)
        border.set(qn("w:color"), "auto")  # Automatic color (black)
        tcBorders.append(border)

    tcPr.append(tcBorders)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0) -> None:
    """Function to set cell margins (padding).

    Args:
        cell: the table cell to modify
        top: margin size in dxa (pt/20). Defaults to 0.
        start: margin size in dxa (pt/20). Defaults to 0.
        bottom: margin size in dxa (pt/20). Defaults to 0.
        end: margin size in dxa (pt/20). Defaults to 0.
    """
    tc = cell._element  # Access the underlying XML element for the cell
    tcPr = tc.find(ns.qn("w:tcPr"))  # Find the <w:tcPr> element if it exists

    # If <w:tcPr> doesn't exist, create it
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    # Create or modify the tcMar (cell margins) element
    tcMar = tcPr.find(ns.qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    # Set each margin (top, start=left, bottom, end=right) in dxa (1/20th of a point)
    for side, margin in [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ]:
        margin_element = OxmlElement(f"w:{side}")
        margin_element.set(ns.qn("w:w"), str(margin))  # Set margin size in dxa
        margin_element.set(ns.qn("w:type"), "dxa")  # dxa = 1/20th of a point
        tcMar.append(margin_element)


def get_primary_resource(
    appointment_id: int,
    relevant_date: datetime,
    api: CTAPI,
    considered_resource_ids: list[int],
) -> str:
    """Helper which is used to get the primary resource allocation of an event.

    Args:
        appointment_id: id of calendar entry
        relevant_date: the date of the appointment is required because appointment_id is not unique on repetitions
        considered_resource_ids: resource ids to consider

    Returns:
        shortened resource representation
    """
    bookings = api.get_bookings(
        resource_ids=considered_resource_ids,
        from_=relevant_date,
        to_=relevant_date,
        appointment_id=appointment_id,
    )
    return {booking["base"]["resource"]["name"] for booking in bookings}


def get_title_name_services(
    calendar_ids: list[int],
    appointment_id: int,
    relevant_date: datetime,
    api: CTAPI,
    considered_program_services: list[int],
    considered_groups: list[int],
) -> str:
    """Helper function which retrieves a text representation of a service including the persons title based on considered groups.

    1. Lookup relevant services
    2. Lookup the prefix of the person to be used based on group assignemnts

    Args:
        calendar_ids: list of calendars to consider
        appointment_id: number of the calendar appointment
        relevant_date: the date of the event to be unique
        api: reference to api in order to request more information from CT
        considered_program_services: list of services which should be considered
        considered_groups: groups which should be used as prefix if applicable

    Returns:
        formatted useable string with title and name
    """
    relevant_event = api.get_event_by_calendar_appointment(
        appointment_id=appointment_id, start_date=relevant_date
    )

    relevant_persons = []
    for service_id in considered_program_services:
        relevant_persons.extend(
            api.get_persons_with_service(
                eventId=relevant_event["id"], serviceId=service_id
            )
        )

    names_with_title = []
    for person in relevant_persons:
        title_prefix = get_group_title_of_person(
            person_id=person["personId"], relevant_groups=considered_groups, api=api
        )
        if person["personId"]:
            lastname = person["person"]["domainAttributes"]["lastName"]
            formatted_name = f"{title_prefix} {lastname}".lstrip()
        else:
            formatted_name = "Noch unbekannt"
        names_with_title.append(formatted_name)

    return ", ".join(names_with_title)


def get_group_title_of_person(
    person_id: int, relevant_groups: list[int], api: CTAPI
) -> str:
    """Retrieve name of first group for specified person and gender if possible.

    Args:
        person_id: CT id of the user
        relevant_groups: the CT id of any groups to be considered as title
        ct_api: access to request more info from CT

    Permissions:
        view person
        view alldata (Persons)
        view group

    Returns:
        Prefix which is used as title incl. gendered version
    """
    for group_id in relevant_groups:
        group_member_ids = [
            group["personId"] for group in api.get_group_members(group_id=group_id)
        ]
        if person_id in group_member_ids:
            group_name = api.get_groups(group_id=group_id)[0]["name"]
            break
    else:
        group_name = ""

    # add 'IN' suffix to first part of group member in order to apply german gender for most common cases
    person = api.get_persons(ids=[person_id])[0]

    gender_map = api.get_persons_masterdata(resultClass="sexes", returnAsDict=True)

    if gender_map[person["sexId"]] == "sex.female" and len(group_name) > 0:
        parts = group_name.split(" ")
        part1_gendered = parts[0] + "in"
        group_name = " ".join([part1_gendered] + parts[1:])

    return group_name


def get_group_name_services(
    calendar_ids: list[int],
    appointment_id: int,
    relevant_date: datetime,
    api: CTAPI,
    considered_music_services: list[int],
    considered_grouptype_role_ids: list[int],
) -> str:
    """Helper which will retrieve the name of special services involved with the calendar appointment on that day.

    1. get event
    2. for each service
        - check who is assigned
        - check what groups are relevant
    3. for each person in service
        - check group membership
        - add groupname to result list if applicable
    4. join the groups applicable to a useable text string

    Args:
        calendar_ids: list of calendars to consider
        appointment_id: number of the calendar appointment
        relevant_date: the date of the event to be unique
        api: reference to api in order to request more information from CT
        considered_music_services: list of services which should be considered
        considered_grouptype_role_ids: list of grouptype_id roles to be considered (differs by group type!)

    Returns:
        text which can be used as suffix - empty in case no special service
    """
    relevant_event = api.get_event_by_calendar_appointment(
        appointment_id=appointment_id, start_date=relevant_date
    )

    result_groups = []
    for service in considered_music_services:
        service_assignments = api.get_persons_with_service(
            eventId=relevant_event["id"], serviceId=service
        )
        persons = [service["person"] for service in service_assignments]
        relevant_group_results = [
            service["groupIds"]
            for service in api.get_event_masterdata()["services"]
            if service["id"] in considered_music_services
        ]
        considered_group_ids = {
            int(group_id)
            for group_result in relevant_group_results
            for group_id in group_result
        }
        for person in persons:
            group_assignemnts = api.get_groups_members(
                group_ids=considered_group_ids,
                grouptype_role_ids=considered_grouptype_role_ids,
                person_ids=[int(person["domainIdentifier"])],
            )
            relevant_group_ids = [group["groupId"] for group in group_assignemnts]
            for group_id in relevant_group_ids:
                group = api.get_groups(group_id=group_id)[0]
                result_groups.append(group["name"])
    return "mit " + " und ".join(result_groups) if len(result_groups) > 0 else ""


def replace_special_services_with_service_shortnames(special_services: str) -> str:
    """Helper which replaces long "special service" strings by their shortform.

    Args:
        special_services: original special service detected
        e.g. "mit Posaunenchor" or "mit Posaunenchor und Kirchenchor"

    Returns:
        short version of special services without mit and separated by , only
    """
    special_services = special_services.split(" und ")
    short_services = []

    for special_service in special_services:
        match special_service.removeprefix("mit "):
            case "Posaunenchor":
                short_services.append("Pos.Chor")
            case "InJoy Chor":
                short_services.append("InJ.Chor")
            case "Kirchenchor":
                short_services.append("Kir.Chor")
            case _:
                short_services.append(special_service.removeprefix("mit "))
                logger.warning("no known shortservice version of %s", special_service)
    return ", ".join(short_services)
